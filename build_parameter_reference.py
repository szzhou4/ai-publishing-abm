"""
build_parameter_reference.py — Rebuild ABM_Agent_Parameter_Reference.docx from scratch.

Version history
---------------
v1  Initial release
v2  Logistic → percentile-threshold acceptance model; piecewise quality penalty
v3  TIER_PRESSURE_TARGETS; tier2_publications; two-component pressure
v4  Resubmission pipeline (MAX_TIER_ATTEMPTS, original_tier, tier_attempts,
    resubmission_queue); Paper.original_tier / tier_attempts fields
v5  Per-paper RL update (fires once per evaluated paper, not once per period);
    Balanced T2 threshold lowered 0.60 → 0.45; Period 0 in figures
v6  T3_FLOOR_FRACTION policy: R1 avoids Tier 3 unless total_pubs < 40% of
    tenure target; applies to new paper assignment and resubmission cascade

Usage:
    python3 build_parameter_reference.py
"""

import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
os.chdir(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand colors ──────────────────────────────────────────────────────────────
PRIMARY   = RGBColor(0x98, 0x1A, 0x31)
SECONDARY = RGBColor(0x00, 0x54, 0x6B)
GOLD      = RGBColor(0xC4, 0xA3, 0x5A)

DOC_PATH = 'ABM_Agent_Parameter_Reference.docx'

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.15)
    sec.right_margin  = Inches(1.15)


# ── Helper utilities ──────────────────────────────────────────────────────────

def heading(text, level=1, color=PRIMARY):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def body(text, italic=False, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    if italic:
        for run in p.runs:
            run.font.italic = True
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    return p

def code(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x20, 0x20, 0x60)
    return p

def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def add_table(headers, rows, header_fill='981A31', header_text_color=None):
    """Add a styled table. header_fill = hex string without #."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light List'
    # Header row
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for par in hdr[i].paragraphs:
            par.paragraph_format.space_after = Pt(0)
            for run in par.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if not header_text_color else header_text_color
    # Data rows — index into pre-allocated rows (do NOT call add_row(); that
    # would append extra rows on top of the ones already created above)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        fill = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            set_cell_shading(cells[ci], fill)
            for par in cells[ci].paragraphs:
                par.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()
    return t


# ═══════════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════════

title = doc.add_heading('ABM Agent & Parameter Reference', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = PRIMARY

sub = doc.add_paragraph('AI Adoption in Academic Publishing — Computational Model')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    run.font.color.rgb = SECONDARY
    run.font.italic = True

ver = doc.add_paragraph('Version 9  |  R1 production 1.5/period · Resubmission cap 2/period · Quality-priority queue')
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in ver.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
body(
    'This document is the authoritative reference for every agent attribute, '
    'model parameter, and core function in the ABM. Edit config.py to change '
    'parameter values; this document should be rebuilt (python3 build_parameter_reference.py) '
    'whenever parameters change.',
    italic=True
)
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# §1  SIMULATION SCALE
# ═══════════════════════════════════════════════════════════════════════════════

heading('§1  Simulation Scale', level=1)
add_table(
    headers=['Parameter', 'Value', 'Notes'],
    rows=[
        ['N_SCHOLARS',     '200',  '25% each institution type'],
        ['N_PERIODS',      '12',   '1 period = 1 semester; 12 periods = 6-year tenure clock'],
        ['N_REPLICATIONS', '50',   'Monte Carlo replications for CI estimation'],
        ['RANDOM_SEED',    '42',   'Seed for first replication; seed+rep for subsequent'],
    ]
)


# ═══════════════════════════════════════════════════════════════════════════════
# §2  SCHOLAR AGENT
# ═══════════════════════════════════════════════════════════════════════════════

heading('§2  Scholar Agent', level=1)
body('File: agents/scholar.py')

heading('2a  Static Attributes (set at initialization, never change)', level=2, color=SECONDARY)
add_table(
    headers=['Attribute', 'Type', 'Range', 'Description'],
    header_fill='00546B',
    rows=[
        ['scholar_id',        'int',   '0 … N−1', 'Unique identifier'],
        ['research_capacity', 'float', '[0.01, 0.99]',
         'Talent + institutional resources. Drawn from Normal(mean, sd) by institution type, '
         'clipped to [0.01, 0.99]. Determines baseline paper quality.'],
        ['institution_type',  'str',   'R1 / R2 / Balanced / Teaching',
         'Determines capacity distribution, base production rate, tier targets, RL pressure.'],
    ]
)

heading('2b  Dynamic Attributes (evolve during simulation)', level=2, color=SECONDARY)
add_table(
    headers=['Attribute', 'Type', 'Initial value', 'Description'],
    header_fill='00546B',
    rows=[
        ['ai_use_level',       'float', 'Drawn from Nag et al. (2025) distribution',
         'Proportion of work done via AI tools. Updated by RL after each paper evaluation. '
         'Clipped to [0, 1].'],
        ['papers_per_period',  'float', 'Computed from base_rate × 2^ai_use_level',
         'Expected new submissions per period (Poisson rate). Recomputed after every RL update.'],
        ['publication_record', 'list',  '[]',
         'List of dicts {period, quality, tier, accepted} for every evaluated paper '
         '(new submissions accepted and rejected; resubmissions recorded only when accepted).'],
        ['resubmission_queue', 'list',  '[]',
         'List of Paper objects awaiting re-evaluation. Papers enter when first rejected; '
         'exit when accepted, downgraded to next tier, or abandoned after MAX_TIER_ATTEMPTS '
         'rejections at Tier 3.'],
    ]
)

heading('2c  Research Capacity Parameters (RESEARCH_CAPACITY_PARAMS)', level=2, color=SECONDARY)
body(
    'research_capacity ~ Normal(mean, sd), clipped to [0.01, 0.99]. '
    'Calibrated so that at zero AI use (quality = research_capacity), scholars produce '
    'papers in realistic quality ranges for each institution type.'
)
add_table(
    headers=['Institution', 'Mean', 'SD', 'Approx. quality range at ai_use = 0'],
    rows=[
        ['R1',       '0.80', '0.05', '0.70 – 0.90  (near T1 threshold ≈ 0.876)'],
        ['R2',       '0.70', '0.10', '0.50 – 0.90  (spans T2 threshold ≈ 0.707)'],
        ['Balanced', '0.60', '0.15', '0.30 – 0.90  (near T2 threshold; most ≥ 0.55 → T2)'],
        ['Teaching', '0.50', '0.15', '0.20 – 0.80  (many T3; strong papers reach T2)'],
    ]
)

heading('2d  Base Production Rates (BASE_PRODUCTION_RATE)', level=2, color=SECONDARY)
body(
    'papers_per_period = BASE_PRODUCTION_RATE[inst] × AI_PRODUCTIVITY_MULTIPLIER ^ ai_use_level\n'
    'AI_PRODUCTIVITY_MULTIPLIER = 2.0 (calibrated to Noy & Zhang 2023; Dell\'Acqua et al. 2023).'
)
add_table(
    headers=['Institution', 'Base rate', 'At ai=0', 'At ai=0.5', 'At ai=1.0'],
    rows=[
        ['R1',       '1.500  ★', '1.5/period', '2.12/period', '3.0/period'],
        ['R2',       '1.500',    '1.5/period', '2.12/period', '3.0/period'],
        ['Balanced', '1.000',    '1.0/period', '1.41/period', '2.0/period'],
        ['Teaching', '0.500',    '0.5/period', '0.71/period', '1.0/period'],
    ]
)
body('★ R1 base rate reduced from 2.0 to 1.5 (v9) to address over-production.', italic=True)

heading('2e  Tenure Targets (TENURE_TARGET_MIDPOINTS)', level=2, color=SECONDARY)
body(
    'v8: Changed from ranges to fixed publication counts. '
    'Tier requirements specify minimum publications in a given tier needed for tenure.'
)
add_table(
    headers=['Institution', 'Total pubs required', 'Tier requirement', 'Notes'],
    rows=[
        ['R1',       '10', '3 Tier 1 pubs',  'Highest bar; T1 pressure drives AI RL'],
        ['R2',        '8', '1 Tier 1 pub',   'Moderate bar; T1 pressure until 1 T1 pub earned'],
        ['Balanced',  '5', '2 Tier 2 pubs',  'No T1 requirement; T2 pressure drives RL'],
        ['Teaching',  '2', 'None',            'No tier-level requirement; overall pressure only'],
    ]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# §3  RL UPDATE RULE
# ═══════════════════════════════════════════════════════════════════════════════

heading('§3  Reinforcement Learning Update Rule', level=1)
body(
    'update_ai_use() is called once per evaluated paper — both new submissions and '
    'resubmissions — rather than once per period. All papers in a period are produced '
    'and evaluated first (so within-period quality calculations use a consistent '
    'ai_use snapshot); RL updates are then applied sequentially.'
)

heading('3a  Update equations', level=2, color=SECONDARY)
code('Acceptance:  ai_use += BASE_STEP_POS × (1 + PRESSURE_WEIGHT × pressure)')
code('Rejection:   ai_use -= BASE_STEP_NEG')
code('ai_use is clipped to [0, 1] after each update.')

add_table(
    headers=['Parameter', 'Value', 'Role'],
    rows=[
        ['BASE_STEP_POS',   '0.05', 'Base positive step per accepted paper'],
        ['BASE_STEP_NEG',   '0.01', 'Negative step per rejected paper (asymmetric; smaller because '
                                    'scholars do not fully abandon AI after a single rejection)'],
        ['PRESSURE_WEIGHT', '1.0',  'Amplification of positive step by publication pressure'],
    ]
)

heading('3b  Pressure calculation', level=2, color=SECONDARY)
body('pressure = max(overall_pressure, tier_pressure)')
code('overall_pressure = max(0, 1 − total_accepted_pubs / TENURE_TARGET_MIDPOINTS[inst])')
code('tier_pressure    = max(0, 1 − tier_pubs / TIER_PRESSURE_TARGETS[inst]["mid"])')
body(
    'tier_pubs is Tier 1 count for R1/R2 and Tier 2 count for Balanced. '
    'Teaching has no tier expectation (tier_pressure = 0).'
)

add_table(
    headers=['Institution', 'Tier tracked', 'Fixed target', 'Notes'],
    rows=[
        ['R1',       'Tier 1', '3', 'Pressure until 3 T1 pubs accumulated (v8: was 4)'],
        ['R2',       'Tier 1', '1', 'Pressure until 1 T1 pub accumulated (v8: was 2)'],
        ['Balanced', 'Tier 2', '2', 'Pressure until 2 T2 pubs accumulated (v8: was 3)'],
        ['Teaching', '—',      '—', 'tier_pressure = 0; overall pressure only'],
    ]
)

heading('3c  Per-paper break-even', level=2, color=SECONDARY)
body(
    'For AI use to grow on average, the per-paper acceptance probability must exceed:'
)
code('p_accept > BASE_STEP_NEG / (BASE_STEP_POS × (1 + PRESSURE_WEIGHT) + BASE_STEP_NEG)')
code('         = 0.01 / (0.05 × 2 + 0.01) ≈ 9%')
body(
    'With per-paper RL, scholars who produce many papers at acceptance rates below 9% '
    'accumulate net-negative updates — AI use self-corrects before quality catastrophically '
    'degrades. This is the key difference from the earlier per-period rule, where '
    '"any acceptance in the period" kept AI growing even at very low per-paper rates. '
    'Impact: R2 and Balanced now equilibrate at lower AI use levels (~0.32–0.54 vs. '
    '0.47–0.63 under the per-period rule), which preserves paper quality and '
    'boosts publication counts toward tenure targets.'
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# §4  QUALITY FUNCTION
# ═══════════════════════════════════════════════════════════════════════════════

heading('§4  Quality Function  (compute_quality)', level=1)
body('File: functions/quality.py')
body(
    'Piecewise penalty calibrated to Gartenberg et al. (2026, Organization Science): '
    '~1.28 SD quality decline at high AI use.'
)

code('quality = research_capacity − penalty(ai_use_level)')
code('')
code('If ai_use_level < AI_LOW_THRESHOLD (0.30):')
code('    penalty = AI_LOW_PENALTY × ai_use_level')
code('Else:')
code('    penalty = AI_LOW_PENALTY × AI_LOW_THRESHOLD')
code('           + AI_HIGH_PENALTY × (ai_use_level − AI_LOW_THRESHOLD) ^ AI_HIGH_EXPONENT')

add_table(
    headers=['Parameter', 'Value', 'Role'],
    rows=[
        ['AI_LOW_THRESHOLD', '0.30', 'Below this AI use: linear (mild) penalty regime'],
        ['AI_LOW_PENALTY',   '0.05', 'Slope of penalty in the linear regime'],
        ['AI_HIGH_PENALTY',  '2.50', 'Coefficient for the nonlinear high-AI penalty'],
        ['AI_HIGH_EXPONENT', '1.50', 'Exponent (> 1 = accelerating penalty above threshold)'],
    ]
)

body(
    'Note: Quality is clipped to [0.01, 0.99] after computation. '
    'At ai_use = 0: quality = research_capacity. '
    'At ai_use = 1: quality ≈ research_capacity − 2.53 (can be negative before clipping).'
)

body('Representative quality values by institution × AI use level:', space_after=2)
add_table(
    headers=['Institution (mean capacity)', 'ai=0', 'ai=0.30', 'ai=0.50', 'ai=0.70'],
    rows=[
        ['R1 (0.85)',       '0.850', '0.835', '0.611', '0.354'],
        ['R2 (0.70)',       '0.700', '0.685', '0.461', '0.204 → clipped'],
        ['Balanced (0.60)', '0.600', '0.585', '0.361 → T3', '0.104 → clipped'],
        ['Teaching (0.50)', '0.500', '0.485', '0.261 → T3', 'near 0 → clipped'],
    ]
)


# ═══════════════════════════════════════════════════════════════════════════════
# §5  JOURNAL TIER ASSIGNMENT
# ═══════════════════════════════════════════════════════════════════════════════

heading('§5  Journal Tier Assignment  (assign_journal_tier)', level=1)
body('File: functions/quality.py  |  Called for new papers only; not for resubmissions.')
body(
    'Assigns a submission tier based on paper quality, institution type, and '
    'whether the scholar has met their Tier 1 publication target (TIER1_PUB_TARGETS). '
    'R1 scholars below their Tier 1 target always submit to Tier 1 (threshold = 0.00), '
    'overriding quality considerations — tenure pressure makes any T1 chance worth taking.'
)

heading('5a  TIER_THRESHOLDS', level=2, color=SECONDARY)
body(
    'v8: All institutions now use the same quality thresholds (T1 ≥ 0.82, T2 ≥ 0.55). '
    'Institution-specific behaviour emerges from differences in research capacity '
    'rather than hard-coded tier exclusions. Only R1 "behind target" retains a special '
    'rule (always attempt T1 under tenure pressure). '
    '"behind" = tier1_pubs < TIER1_PUB_TARGET; "ahead" = target met.'
)
add_table(
    headers=['Institution', 'State', 'T1 min quality', 'T2 min quality', 'Notes'],
    rows=[
        ['R1', 'behind T1 target', '0.00',
         '0.00 (→ T1)', 'Always T1 until 3 T1 pubs earned; tenure pressure overrides quality'],
        ['R1', 'ahead of T1 target', '0.82', '0.55',
         'Quality-based targeting; same thresholds as all other institutions'],
        ['R2',       'behind / ahead', '0.82', '0.55',
         'Unified threshold; T1 attempts rare but possible for strong papers'],
        ['Balanced', 'behind / ahead', '0.82', '0.55',
         'Unified threshold; T1 very rare given capacity 0.60; T2 achievable'],
        ['Teaching', 'behind / ahead', '0.82', '0.55',
         'Unified threshold; T2 possible for exceptional papers (capacity ≥ 0.55)'],
    ]
)

heading('5b  TIER1_PUB_TARGETS', level=2, color=SECONDARY)
body(
    'Number of Tier 1 publications at which R1/R2 switch from pressure-based T1 targeting '
    'to quality-based targeting (has no effect for Balanced/Teaching, target = 0).'
)
add_table(
    headers=['Institution', 'TIER1_PUB_TARGET', 'Notes'],
    rows=[
        ['R1',       '3', 'v8: fixed T1 requirement (was 4); switch to quality-based once met'],
        ['R2',       '1', 'v8: fixed T1 requirement (was 2); switch to quality-based once met'],
        ['Balanced', '0', 'No Tier 1 requirement; quality-based throughout'],
        ['Teaching', '0', 'No Tier 1 requirement'],
    ]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# §6  ACCEPTANCE MODEL
# ═══════════════════════════════════════════════════════════════════════════════

heading('§6  Acceptance Model  (compute_acceptance_probability)', level=1)
body('File: functions/quality.py  |  Called for every paper evaluation including resubmissions.')
body(
    'Percentile-threshold model: a journal with X% acceptance rate accepts the top X% '
    'of papers by quality. Thresholds derived from the assumed population quality '
    'distribution Normal(QUALITY_DIST_MEAN=0.50, QUALITY_DIST_SD=0.20).'
)

code('threshold_tier = ppf(1 − base_accept_prob,  mean=0.50,  sd=0.20)')
code('')
code('noisy_quality  = clip(quality + Normal(0, QUALITY_NOISE_SD),  0.01,  0.99)')
code('accepted       = noisy_quality >= threshold_tier')
code('')
code('accept_prob    = Φ((quality − threshold_tier) / QUALITY_NOISE_SD)   [analytical]')

add_table(
    headers=['Tier', 'Base accept rate', 'Threshold', 'Notes'],
    rows=[
        ['Tier 1 (Top)',   '3%',  '≈ 0.876', 'ppf(0.97, 0.50, 0.20)  v8: tightened from 5% / 0.829'],
        ['Tier 2 (Mid)',   '15%', '≈ 0.707', 'ppf(0.85, 0.50, 0.20)'],
        ['Tier 3 (Lower)', '30%', '≈ 0.605', 'ppf(0.70, 0.50, 0.20)'],
    ]
)

add_table(
    headers=['Parameter', 'Value', 'Role'],
    rows=[
        ['QUALITY_DIST_MEAN', '0.50', 'Assumed mean of the population quality distribution'],
        ['QUALITY_DIST_SD',   '0.20', 'Assumed SD of the population quality distribution'],
        ['QUALITY_NOISE_SD',  '0.10', 'Gaussian noise added to quality at review '
                                      '(captures reviewer randomness; new draw each evaluation)'],
    ]
)


# ═══════════════════════════════════════════════════════════════════════════════
# §7  RESUBMISSION PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════

heading('§7  Tier 3 Submission / Cascade Policy', level=1)
body(
    'R1 scholars have strong institutional norms against Tier 3 publications. '
    'The T3_FLOOR_FRACTION policy gates T3 submission in two places:'
)
for b in [
    'New paper tier assignment: if assign_journal_tier() returns Tier 3 for an R1 scholar '
    'who is NOT desperate, the tier is redirected to Tier 2 instead.',
    'Resubmission cascade: when an R1 paper would downgrade from Tier 2 to Tier 3 after '
    'MAX_TIER_ATTEMPTS rejections, it is ABANDONED rather than downgraded, unless the '
    'scholar is desperate.',
    '"Desperate" = total_publications < T3_FLOOR_FRACTION[institution] × TENURE_TARGET_MIDPOINTS[institution].',
    'For institutions with T3_FLOOR_FRACTION = None, Tier 3 is unrestricted (R2, Balanced, Teaching).',
]:
    bullet(b)

add_table(
    headers=['Institution', 'T3_FLOOR_FRACTION', 'Desperation threshold', 'Notes'],
    rows=[
        ['R1',       '0.40', '< 4 total pubs  (40% × 10)',
         'T3 only if well behind overall target; papers otherwise redirected/abandoned'],
        ['R2',       'None', '—', 'Unrestricted; T3 cascade always allowed'],
        ['Balanced', 'None', '—', 'Unrestricted'],
        ['Teaching', 'None', '—', 'Always Tier 3 anyway'],
    ]
)
body(
    'Result (v6 simulation): R1 T3 publications collapsed from 4.8 to 0.2 per scholar. '
    'Papers that previously cascaded to T3 are now redirected to T2 on initial submission '
    'or abandoned after 3 T2 rejections, consistent with R1 institutional norms. '
    'Redirected papers generate additional T2 rejections, providing corrective RL signals '
    'that modestly lower R1 final AI use (0.554 → 0.500).',
    italic=True
)
doc.add_page_break()

heading('§9  Resubmission Pipeline', level=1)

heading('7a  Design principles', level=2, color=SECONDARY)
for b in [
    'Rejected papers enter scholar.resubmission_queue at end of the period in which '
    'they were first rejected (tier_attempts set to 1 before queuing).',
    'Each subsequent period, queued papers are evaluated BEFORE new paper production. '
    'Same paper quality is used; a new Gaussian noise draw is applied (QUALITY_NOISE_SD = 0.10).',
    'After MAX_TIER_ATTEMPTS (= 3) successive rejections at the current tier: if journal_tier < 3, '
    'downgrade (journal_tier += 1, tier_attempts = 0); if at Tier 3, abandon the paper.',
    'Accepted resubmissions are recorded in publication_record and count toward '
    'total_publications, tier-specific counts, and the RL signal. They do NOT count '
    'toward papers_per_period (efficiency metric).',
    'Rejected resubmissions trigger a negative RL update (per-paper RL rule) but are NOT '
    'recorded in publication_record (avoids inflating the total_produced denominator).',
]:
    bullet(b)

heading('7b  Paper fields added for resubmission tracking', level=2, color=SECONDARY)
add_table(
    headers=['Field', 'Type', 'Default', 'Description'],
    header_fill='00546B',
    rows=[
        ['original_tier',  'int', '0',
         'Tier at first submission; 0 = not yet assigned. Set at paper production time.'],
        ['tier_attempts',  'int', '0',
         'Successive rejections at the current journal_tier. Reset to 0 on tier downgrade. '
         'Set to 1 immediately after new-paper rejection (before queuing).'],
    ]
)

heading('7c  Pipeline parameters', level=2, color=SECONDARY)
add_table(
    headers=['Parameter', 'Value', 'Description'],
    rows=[
        ['MAX_TIER_ATTEMPTS', '3',
         'Maximum rejections at one tier before downgrading. '
         'Paper gets up to 3 tries at T1, then 3 at T2, then 3 at T3 (9 periods max).'],
        ['MAX_RESUB_PER_PERIOD', '2',
         'Maximum resubmissions evaluated per scholar per period (v9). '
         'If more papers are queued, the top-quality papers are evaluated first; '
         'lower-quality papers are deferred to the next period with no evaluation, '
         'no RL signal, and no change to tier_attempts.'],
    ]
)

heading('7d  Example trajectory', level=2, color=SECONDARY)
body('Paper produced in period 1, originally T1:', space_after=2)
add_table(
    headers=['Period', 'Tier', 'tier_attempts', 'Outcome', 'Action'],
    rows=[
        ['1', 'T1', '0→1', 'Rejected', 'Set tier_attempts=1, add to queue'],
        ['2', 'T1', '1→2', 'Rejected', 'Increment; stays in queue'],
        ['3', 'T1', '2→3', 'Rejected', '3 = MAX_TIER_ATTEMPTS → downgrade to T2, reset=0'],
        ['4', 'T2', '0→1', 'Rejected', 'Increment; stays in queue'],
        ['5', 'T2', '1→2', 'Accepted','Record as T2 publication; remove from queue'],
        ['—', '—', '—', '—', 'If rejected at T2 × 3 → downgrade to T3; if T3 × 3 → abandon'],
    ]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# §8  QUICK REFERENCE — ALL PARAMETERS
# ═══════════════════════════════════════════════════════════════════════════════

heading('§10  Quick Reference — All Parameters (config.py)', level=1)

add_table(
    headers=['Parameter', 'Value', 'Section'],
    rows=[
        # Simulation scale
        ['N_SCHOLARS',             '200',   '§1'],
        ['N_PERIODS',              '12',    '§1'],
        ['N_REPLICATIONS',         '50',    '§1'],
        ['RANDOM_SEED',            '42',    '§1'],
        # Institution distribution
        ['INSTITUTION_DISTRIBUTION', '[0.25, 0.25, 0.25, 0.25]', '§2'],
        # Research capacity
        ['R1 capacity',      'N(0.80, 0.05) clip [0.01,0.99]', '§2c  v8: mean 0.85→0.80'],
        ['R2 capacity',      'N(0.70, 0.10) clip [0.01,0.99]', '§2c'],
        ['Balanced capacity','N(0.60, 0.15) clip [0.01,0.99]', '§2c'],
        ['Teaching capacity','N(0.50, 0.15) clip [0.01,0.99]', '§2c'],
        # Production rates
        ['BASE_PRODUCTION_RATE R1',       '2.000', '§2d'],
        ['BASE_PRODUCTION_RATE R2',       '1.500', '§2d'],
        ['BASE_PRODUCTION_RATE Balanced', '1.000', '§2d'],
        ['BASE_PRODUCTION_RATE Teaching', '0.500', '§2d'],
        ['AI_PRODUCTIVITY_MULTIPLIER',    '2.0',   '§2d'],
        # Tenure targets (fixed; v8)
        ['TENURE_TARGET R1',       '10  (+ 3 T1 pubs)', '§2e'],
        ['TENURE_TARGET R2',        '8  (+ 1 T1 pub)',  '§2e  v8: was 9'],
        ['TENURE_TARGET Balanced',  '5  (+ 2 T2 pubs)', '§2e  v8: was 6'],
        ['TENURE_TARGET Teaching',  '2  (no tier req.)', '§2e'],
        # RL
        ['BASE_STEP_POS',   '0.05', '§3a  — per accepted paper'],
        ['BASE_STEP_NEG',   '0.01', '§3a  — per rejected paper'],
        ['PRESSURE_WEIGHT', '1.0',  '§3a'],
        ['RL firing',       'Per paper (new + resubmission)', '§3  ★ changed v5→v6'],
        # Quality function
        ['AI_LOW_THRESHOLD', '0.30', '§4'],
        ['AI_LOW_PENALTY',   '0.05', '§4'],
        ['AI_HIGH_PENALTY',  '2.50', '§4'],
        ['AI_HIGH_EXPONENT', '1.50', '§4'],
        # Tier thresholds (v8: unified across all institutions)
        ['T1 min quality — R1 (behind)',  '0.00',  '§5a  pressure override: always T1'],
        ['T1 min quality — all others',   '0.82',  '§5a  v8: unified threshold'],
        ['T2 min quality — all',          '0.55',  '§5a  v8: unified threshold'],
        # Tier 1 pub targets
        ['TIER1_PUB_TARGETS R1',       '3', '§5b  v8: was 4'],
        ['TIER1_PUB_TARGETS R2',       '1', '§5b  v8: was 2'],
        ['TIER1_PUB_TARGETS Balanced', '0', '§5b'],
        ['TIER1_PUB_TARGETS Teaching', '0', '§5b'],
        # Acceptance model
        ['QUALITY_DIST_MEAN',  '0.50',   '§6'],
        ['QUALITY_DIST_SD',    '0.20',   '§6'],
        ['QUALITY_NOISE_SD',   '0.10',   '§6'],
        ['T1 base accept rate','3%  → threshold ≈ 0.876', '§6  v8: was 5% / 0.829'],
        ['T2 base accept rate','15% → threshold ≈ 0.707', '§6'],
        ['T3 base accept rate','30% → threshold ≈ 0.605', '§6'],
        # T3 avoidance
        ['T3_FLOOR_FRACTION R1',       '0.40 ★',
         '§7  ★ added v6 — T3 only if total_pubs < 4'],
        ['T3_FLOOR_FRACTION R2/Bal/Teach', 'None', '§7  — unrestricted'],
        # Resubmission
        ['MAX_TIER_ATTEMPTS',    '3', '§7c'],
        ['MAX_RESUB_PER_PERIOD', '2', '§7c  v9: top-quality papers evaluated first; rest deferred'],
        ['BASE_PRODUCTION_RATE R1', '1.500', '§2d  v9: reduced from 2.000'],
        # Initial AI use
        ['NAG_AI_VALUES',      '[0.00, 0.05, 0.15, 0.30, 0.50, 0.70]', 'Nag et al. (2025)'],
        ['NAG_AI_PROPORTIONS', '[0.145, 0.177, 0.048, 0.177, 0.258, 0.194]',
         'Nag et al. (2025); same distribution across all institution types'],
    ]
)

body(
    'v8 marks parameters changed in Version 8 (current).',
    italic=True, space_after=3
)


# ═══════════════════════════════════════════════════════════════════════════════
# §9  VERSION HISTORY
# ═══════════════════════════════════════════════════════════════════════════════

heading('§11  Version History', level=1)
add_table(
    headers=['Version', 'Key changes'],
    rows=[
        ['v1', 'Initial release. Logistic acceptance function. Per-period RL.'],
        ['v2', 'Percentile-threshold acceptance model. Piecewise quality penalty '
               '(Gartenberg et al. 2026). Quality noise SD = 0.10.'],
        ['v3', 'TIER_PRESSURE_TARGETS added. tier2_publications tracked. '
               'Two-component pressure: max(overall, tier). '
               'R2 TIER1_PUB_TARGET = 2 added.'],
        ['v4', 'Resubmission pipeline: MAX_TIER_ATTEMPTS = 3; Paper.original_tier / '
               'tier_attempts fields; Scholar.resubmission_queue.'],
        ['v5', 'Per-paper RL firing (was per-period). '
               'Balanced T2 threshold 0.60 → 0.45. '
               'Period 0 recorded in sim_results.csv for figures.'],
        ['v6', 'T3_FLOOR_FRACTION policy: R1 avoids Tier 3 unless total_pubs < 40% '
               'of tenure target (< 4 pubs). Applies to both new paper tier assignment '
               '(redirected T3→T2 if not desperate) and resubmission cascade (abandoned '
               'rather than downgraded if not desperate). R1 T3 pubs: 4.8 → 0.2.'],
        ['v7', 'R1 mean capacity 0.85 → 0.80. Unified tier thresholds: T1 ≥ 0.82, T2 ≥ 0.55 '
               'for all institutions (was institution-specific with hard 9.99 blocks for '
               'Teaching/Balanced). Teaching and Balanced can now submit to higher tiers '
               'when paper quality warrants it.'],
        ['v8', 'R1 SD 0.10 → 0.05. Fixed tenure targets: R1=10, R2=8, Balanced=5, Teaching=2 '
               '(replaced ranges). Fixed tier requirements: R1 needs 3 T1 pubs, R2 needs 1 T1 pub, '
               'Balanced needs 2 T2 pubs, Teaching has no tier requirement. '
               'T1 base acceptance rate tightened 5% → 3% (threshold ≈ 0.829 → 0.876). '
               'TIER_PRESSURE_TARGETS updated to match fixed requirements.'],
        ['v9', 'R1 base production rate 2.0 → 1.5 papers/period (R1 was over-producing). '
               'MAX_RESUB_PER_PERIOD = 2: at most 2 resubmissions evaluated per scholar per period; '
               'excess papers sorted by quality descending and deferred to the next period '
               'with no evaluation, no RL signal, and no tier_attempts increment.'],
    ]
)

doc.save(DOC_PATH)
print(f'Saved {DOC_PATH}')
