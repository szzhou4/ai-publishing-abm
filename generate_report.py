"""
generate_report.py — Regenerate report figures and ABM_Simulation_Results.docx (v4).
Run after run_simulation.py.

Usage:
    python3 generate_report.py
"""

import sys, os
sys.path.insert(0, os.path.dirname(__file__))
os.chdir(os.path.dirname(os.path.abspath(__file__)))

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import warnings
warnings.filterwarnings('ignore')

from config import (
    INSTITUTION_TYPES, INST_COLORS, COLORS,
    TENURE_TARGET_MIDPOINTS, N_PERIODS,
    TIER1_PUB_TARGETS, TIER_PRESSURE_TARGETS,
)

os.makedirs('outputs/report_figs', exist_ok=True)

# ── Load outputs ──────────────────────────────────────────────────────────────
period_df = pd.read_csv('outputs/sim_results.csv')
end_df    = pd.read_csv('outputs/scholar_endstate.csv')

inst_labels = {
    'R1':       'R1',
    'R2':       'R2',
    'Balanced': 'Balanced',
    'Teaching': 'Teaching-Focused',
}

plt.rcParams.update({
    'font.family':         'sans-serif',
    'axes.spines.top':     False,
    'axes.spines.right':   False,
    'figure.dpi':          150,
    'axes.labelsize':      11,
    'axes.titlesize':      12,
    'legend.fontsize':     9,
    'xtick.labelsize':     9,
    'ytick.labelsize':     9,
})

PERIODS = sorted(period_df['period'].unique())


def ci(df, col):
    """Return mean, lower CI, upper CI across replications by period."""
    g  = df.groupby('period')[col]
    m  = g.mean()
    se = g.sem()
    return m, m - 1.96*se, m + 1.96*se


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — Basic Descriptives
# ═══════════════════════════════════════════════════════════════════════════════

# Fig 1a — AI Use Over Time
fig, ax = plt.subplots(figsize=(8, 4.5))
for inst in INSTITUTION_TYPES:
    m, lo, hi = ci(period_df, f'ai_{inst}')
    c = INST_COLORS[inst]
    ax.plot(PERIODS, m, color=c, lw=2.2, label=inst_labels[inst])
    ax.fill_between(PERIODS, lo, hi, color=c, alpha=0.15)
ax.set_title('AI Use Level Over Time by Institution Type',
             fontweight='bold', color=COLORS['primary'], pad=10)
ax.set_xlabel('Period (0 = initial state; 1–12 = post-RL)')
ax.set_ylabel('Mean AI Use Level')
ax.set_ylim(0, 1); ax.set_xlim(0, N_PERIODS)
ax.axvline(0.5, color='gray', lw=0.8, ls=':', alpha=0.4)
ax.legend(frameon=False, loc='upper left')
plt.tight_layout()
plt.savefig('outputs/report_figs/fig1_ai_use_time.png', bbox_inches='tight')
plt.close()
print("Saved fig1_ai_use_time.png")

# Fig 1b — Acceptance Rate Over Time (new papers + combined with resubmissions)
fig, ax = plt.subplots(figsize=(8, 4.5))
m, lo, hi   = ci(period_df.dropna(subset=['acceptance_rate']), 'acceptance_rate')
m2, lo2, hi2 = ci(period_df.dropna(subset=['acceptance_rate_all']), 'acceptance_rate_all')
ax.plot(m.index,  m*100,  color=COLORS['primary'],   lw=2.2, label='New submissions only')
ax.fill_between(m.index,  lo*100,  hi*100,  color=COLORS['primary'],   alpha=0.15)
ax.plot(m2.index, m2*100, color=COLORS['secondary'], lw=2.2, ls='--',
        label='All evaluations (incl. resubmissions)')
ax.fill_between(m2.index, lo2*100, hi2*100, color=COLORS['secondary'], alpha=0.10)
ax.set_title('Overall Acceptance Rate Over Time',
             fontweight='bold', color=COLORS['primary'], pad=10)
ax.set_xlabel('Period (semester)')
ax.set_ylabel('Acceptance Rate (%)')
ax.set_xlim(0, N_PERIODS); ax.legend(frameon=False)
plt.tight_layout()
plt.savefig('outputs/report_figs/fig1_acceptance_time.png', bbox_inches='tight')
plt.close()
print("Saved fig1_acceptance_time.png")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Efficiency
# ═══════════════════════════════════════════════════════════════════════════════

# Fig 2a — Mean accepted publications per period vs AI use (end-state scatter)
fig, ax = plt.subplots(figsize=(8, 5))
# Compute per-scholar average AI use and publications per period across reps
scholar_summary = (
    end_df.groupby(['institution_type', 'scholar_id'])
    .agg(ai_use=('ai_use_final','mean'), total_pubs=('total_pubs','mean'))
    .reset_index()
)
# Bin AI use into deciles
scholar_summary['ai_bin'] = pd.cut(scholar_summary['ai_use'], bins=10, labels=False)
binned = scholar_summary.groupby(['institution_type','ai_bin']).agg(
    mean_ai=('ai_use','mean'),
    mean_pubs=('total_pubs','mean'),
    sem_pubs=('total_pubs','sem'),
).reset_index().dropna()

for inst in INSTITUTION_TYPES:
    sub = binned[binned['institution_type'] == inst].sort_values('mean_ai')
    c   = INST_COLORS[inst]
    ax.plot(sub['mean_ai'], sub['mean_pubs'] / N_PERIODS,
            color=c, lw=2.0, marker='o', ms=5, label=inst_labels[inst])
    ax.fill_between(sub['mean_ai'],
                    (sub['mean_pubs'] - 1.96*sub['sem_pubs']) / N_PERIODS,
                    (sub['mean_pubs'] + 1.96*sub['sem_pubs']) / N_PERIODS,
                    color=c, alpha=0.12)
ax.set_xlabel('Final AI Use Level')
ax.set_ylabel('Mean Accepted Publications per Period')
ax.set_title('Efficiency: Publications per Period vs. AI Use Level',
             fontweight='bold', color=COLORS['primary'], pad=10)
ax.legend(frameon=False, loc='upper left')
plt.tight_layout()
plt.savefig('outputs/report_figs/fig2_efficiency.png', bbox_inches='tight')
plt.close()
print("Saved fig2_efficiency.png")

# Fig 2b — Acceptance rate (new papers) vs AI use (end-state scatter)
fig, ax = plt.subplots(figsize=(8, 5))
scholar_ar = (
    end_df.groupby(['institution_type', 'scholar_id'])
    .agg(ai_use=('ai_use_final','mean'), acc_rate=('acceptance_rate','mean'))
    .reset_index()
)
scholar_ar['ai_bin'] = pd.cut(scholar_ar['ai_use'], bins=10, labels=False)
binned_ar = scholar_ar.groupby(['institution_type','ai_bin']).agg(
    mean_ai=('ai_use','mean'),
    mean_ar=('acc_rate','mean'),
    sem_ar=('acc_rate','sem'),
).reset_index().dropna()

for inst in INSTITUTION_TYPES:
    sub = binned_ar[binned_ar['institution_type'] == inst].sort_values('mean_ai')
    c   = INST_COLORS[inst]
    ax.plot(sub['mean_ai'], sub['mean_ar']*100,
            color=c, lw=2.0, marker='o', ms=5, label=inst_labels[inst])
    ax.fill_between(sub['mean_ai'],
                    (sub['mean_ar'] - 1.96*sub['sem_ar'])*100,
                    (sub['mean_ar'] + 1.96*sub['sem_ar'])*100,
                    color=c, alpha=0.12)
ax.set_xlabel('Final AI Use Level')
ax.set_ylabel('Acceptance Rate (%)')
ax.set_title('Per-Scholar Acceptance Rate vs. Final AI Use Level',
             fontweight='bold', color=COLORS['primary'], pad=10)
ax.legend(frameon=False)
plt.tight_layout()
plt.savefig('outputs/report_figs/fig2_accept_by_ai.png', bbox_inches='tight')
plt.close()
print("Saved fig2_accept_by_ai.png")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — Quality
# ═══════════════════════════════════════════════════════════════════════════════

# Fig 3a — Mean paper quality over time
fig, ax = plt.subplots(figsize=(8, 4.5))
for inst in INSTITUTION_TYPES:
    col = f'q_{inst}'
    m, lo, hi = ci(period_df.dropna(subset=[col]), col)
    if m.empty: continue
    c = INST_COLORS[inst]
    ax.plot(m.index, m, color=c, lw=2.2, label=inst_labels[inst])
    ax.fill_between(m.index, lo, hi, color=c, alpha=0.15)
# Reference lines for acceptance thresholds
ax.axhline(0.829, color='gray', lw=1.0, ls='--', alpha=0.5, label='T1 threshold (0.829)')
ax.axhline(0.707, color='gray', lw=1.0, ls=':', alpha=0.5,  label='T2 threshold (0.707)')
ax.set_title('Mean New-Paper Quality Over Time by Institution Type',
             fontweight='bold', color=COLORS['primary'], pad=10)
ax.set_xlabel('Period (semester)')
ax.set_ylabel('Mean Paper Quality')
ax.set_xlim(0, N_PERIODS)
ax.legend(frameon=False, loc='lower left', ncol=2)
plt.tight_layout()
plt.savefig('outputs/report_figs/fig3_quality_time.png', bbox_inches='tight')
plt.close()
print("Saved fig3_quality_time.png")

# Fig 3b — Mean quality at acceptance vs AI use (end-state)
fig, ax = plt.subplots(figsize=(8, 5))
scholar_q = (
    end_df.groupby(['institution_type', 'scholar_id'])
    .agg(ai_use=('ai_use_final','mean'),
         q_all=('mean_quality_all','mean'),
         q_acc=('mean_quality_accepted','mean'))
    .reset_index()
)
scholar_q['ai_bin'] = pd.cut(scholar_q['ai_use'], bins=10, labels=False)
binned_q = scholar_q.groupby(['institution_type','ai_bin']).agg(
    mean_ai=('ai_use','mean'),
    mean_q_all=('q_all','mean'),
    mean_q_acc=('q_acc','mean'),
).reset_index().dropna()

for inst in INSTITUTION_TYPES:
    sub = binned_q[binned_q['institution_type'] == inst].sort_values('mean_ai')
    c   = INST_COLORS[inst]
    ax.plot(sub['mean_ai'], sub['mean_q_all'],
            color=c, lw=2.0, ls='-', marker='o', ms=5, label=f'{inst_labels[inst]} (all)')
    ax.plot(sub['mean_ai'], sub['mean_q_acc'],
            color=c, lw=1.4, ls='--', label=f'{inst_labels[inst]} (accepted)')
ax.set_xlabel('Final AI Use Level')
ax.set_ylabel('Mean Paper Quality')
ax.set_title('Paper Quality vs. AI Use Level\n(solid = all papers; dashed = accepted only)',
             fontweight='bold', color=COLORS['primary'], pad=10)
ax.legend(frameon=False, ncol=2, fontsize=8)
plt.tight_layout()
plt.savefig('outputs/report_figs/fig3_quality_vs_ai.png', bbox_inches='tight')
plt.close()
print("Saved fig3_quality_vs_ai.png")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — Tenure Target Performance
# ═══════════════════════════════════════════════════════════════════════════════

# Fig 4a — Total publications vs tenure target (grouped bar)
fig, ax = plt.subplots(figsize=(7, 4.5))
means = end_df.groupby('institution_type')['total_pubs'].mean().reindex(INSTITUTION_TYPES)
sems  = end_df.groupby('institution_type')['total_pubs'].sem().reindex(INSTITUTION_TYPES)
x     = np.arange(len(INSTITUTION_TYPES))
bars  = ax.bar(x, means, yerr=1.96*sems,
               color=[INST_COLORS[i] for i in INSTITUTION_TYPES],
               edgecolor='white', capsize=5, error_kw={'linewidth':1.5})
for xi, inst in enumerate(INSTITUTION_TYPES):
    mid = TENURE_TARGET_MIDPOINTS[inst]
    ax.plot([xi-0.4, xi+0.4], [mid, mid], color='black', lw=2, ls='--')
    ax.text(xi, mid + 0.2, f'Target: {mid}', ha='center', va='bottom',
            fontsize=8, color='black')
ax.set_xticks(x)
ax.set_xticklabels([inst_labels[i] for i in INSTITUTION_TYPES])
ax.set_title('Mean Total Publications vs. Tenure Targets\n(dashed = target midpoint)',
             fontweight='bold', color=COLORS['primary'], pad=10)
ax.set_ylabel('Mean Publications (6-year tenure clock)')
plt.tight_layout()
plt.savefig('outputs/report_figs/fig4_tenure_pubs.png', bbox_inches='tight')
plt.close()
print("Saved fig4_tenure_pubs.png")

# Fig 4b — Tier breakdown by institution type (stacked bar)
fig, ax = plt.subplots(figsize=(7, 4.5))
t1 = end_df.groupby('institution_type')['tier1_pubs'].mean().reindex(INSTITUTION_TYPES)
t2 = end_df.groupby('institution_type')['tier2_pubs'].mean().reindex(INSTITUTION_TYPES)
t3 = end_df.groupby('institution_type')['tier3_pubs'].mean().reindex(INSTITUTION_TYPES)
x  = np.arange(len(INSTITUTION_TYPES))
ax.bar(x, t1, label='Tier 1',  color=COLORS['primary'],   edgecolor='white', alpha=0.9)
ax.bar(x, t2, bottom=t1, label='Tier 2', color=COLORS['secondary'],  edgecolor='white', alpha=0.9)
ax.bar(x, t3, bottom=t1+t2, label='Tier 3', color=COLORS['gold'], edgecolor='white', alpha=0.9)
# Tier targets for R1 and R2
for xi, inst in enumerate(INSTITUTION_TYPES):
    target = TIER1_PUB_TARGETS[inst]
    if target > 0:
        ax.plot([xi-0.4, xi+0.4], [target, target],
                color='black', lw=2, ls='--', zorder=5)
        ax.text(xi, target + 0.15, f'T1 target: {target}', ha='center', va='bottom',
                fontsize=7.5, color='black')
ax.set_xticks(x)
ax.set_xticklabels([inst_labels[i] for i in INSTITUTION_TYPES])
ax.set_title('Mean Publications by Tier and Institution Type',
             fontweight='bold', color=COLORS['primary'], pad=10)
ax.set_ylabel('Mean Publications')
ax.legend(frameon=False)
plt.tight_layout()
plt.savefig('outputs/report_figs/fig4_tier_breakdown.png', bbox_inches='tight')
plt.close()
print("Saved fig4_tier_breakdown.png")

# Fig 4c — % Meeting tenure targets by AI use quartile
fig, ax = plt.subplots(figsize=(8, 5))
end_df['ai_quartile'] = pd.qcut(end_df['ai_use_final'], q=4,
                                labels=['Q1\n(Low AI)', 'Q2', 'Q3', 'Q4\n(High AI)'])
for inst in INSTITUTION_TYPES:
    tgt = TENURE_TARGET_MIDPOINTS[inst]
    sub = end_df[end_df['institution_type'] == inst].copy()
    sub['met'] = sub['total_pubs'] >= tgt
    pct = sub.groupby('ai_quartile', observed=True)['met'].mean() * 100
    ax.plot(pct.index, pct.values, color=INST_COLORS[inst],
            lw=2.0, marker='o', ms=6, label=inst_labels[inst])
ax.set_xlabel('AI Use Quartile')
ax.set_ylabel('% Scholars Meeting Tenure Publication Target')
ax.set_title('Tenure Target Achievement by AI Use Quartile',
             fontweight='bold', color=COLORS['primary'], pad=10)
ax.set_ylim(0, 105)
ax.legend(frameon=False)
plt.tight_layout()
plt.savefig('outputs/report_figs/fig4_tenure_pct_ai.png', bbox_inches='tight')
plt.close()
print("Saved fig4_tenure_pct_ai.png")

# Fig 4d — Tier 1 (or Tier 2 for Balanced) target achievement by AI use quartile
fig, ax = plt.subplots(figsize=(8, 5))
tier_targets = {
    'R1':       ('tier1_pubs', TIER1_PUB_TARGETS['R1'],       'Tier 1 target'),
    'R2':       ('tier1_pubs', TIER1_PUB_TARGETS['R2'],       'Tier 1 target'),
    'Balanced': ('tier2_pubs', TIER_PRESSURE_TARGETS['Balanced']['mid'], 'Tier 2 target'),
    'Teaching': ('tier3_pubs', TENURE_TARGET_MIDPOINTS['Teaching'],      'Any-tier target'),
}
for inst in INSTITUTION_TYPES:
    col, tgt, label = tier_targets[inst]
    if tgt == 0: continue
    sub = end_df[end_df['institution_type'] == inst].copy()
    sub['met'] = sub[col] >= tgt
    pct = sub.groupby('ai_quartile', observed=True)['met'].mean() * 100
    ax.plot(pct.index, pct.values, color=INST_COLORS[inst],
            lw=2.0, marker='o', ms=6, label=f'{inst_labels[inst]} ({label}={tgt})')
ax.set_xlabel('AI Use Quartile')
ax.set_ylabel('% Scholars Meeting Tier-Specific Target')
ax.set_title('Tier-Specific Target Achievement by AI Use Quartile\n'
             '(R1/R2 = Tier 1 target; Balanced = Tier 2 target)',
             fontweight='bold', color=COLORS['primary'], pad=10)
ax.set_ylim(0, 105)
ax.legend(frameon=False, fontsize=8)
plt.tight_layout()
plt.savefig('outputs/report_figs/fig4_tier_target_by_ai.png', bbox_inches='tight')
plt.close()
print("Saved fig4_tier_target_by_ai.png")

print("\nAll figures saved. Building Word doc...")

# ═══════════════════════════════════════════════════════════════════════════════
# BUILD WORD DOC
# ═══════════════════════════════════════════════════════════════════════════════
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY   = RGBColor(0x98, 0x1A, 0x31)
SECONDARY = RGBColor(0x00, 0x54, 0x6B)

DOC_PATH = 'ABM_Simulation_Results.docx'

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.15)
    sec.right_margin  = Inches(1.15)

# ── Styles helper ─────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1, color=PRIMARY):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_figure(doc, path, caption, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.italic = True
            run.font.size   = Pt(9)
        cap.paragraph_format.space_after = Pt(12)

def add_stat_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for par in hdr_cells[i].paragraphs:
            for run in par.runs:
                run.font.bold = True
    # Index into pre-allocated rows — do NOT call add_row() or extra empty
    # rows will appear before the data rows in the rendered document.
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
    doc.add_paragraph()

# ── Cover ─────────────────────────────────────────────────────────────────────
title = doc.add_heading('AI Publishing ABM — Simulation Results', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = PRIMARY

sub = doc.add_paragraph('Agent-Based Model of AI Adoption in Academic Publishing  |  Version 3')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(4)
for run in sub.runs:
    run.font.color.rgb = SECONDARY

sub2 = doc.add_paragraph(
    f'N = 200 scholars  ·  12 periods (6-year tenure clock)  ·  50 replications\n'
    'Resubmission pipeline active: MAX_TIER_ATTEMPTS = 3 per tier')
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub2.runs:
    run.font.size = Pt(10)
    run.font.italic = True

doc.add_paragraph()

# ── Key stats callout table ───────────────────────────────────────────────────
add_heading(doc, 'Key Simulation Statistics at a Glance', level=1, color=PRIMARY)

# Compute summary numbers
mean_ai_final = end_df.groupby('institution_type')['ai_use_final'].mean()
mean_pubs      = end_df.groupby('institution_type')['total_pubs'].mean()
mean_t1        = end_df.groupby('institution_type')['tier1_pubs'].mean()
mean_t2        = end_df.groupby('institution_type')['tier2_pubs'].mean()
mean_t3        = end_df.groupby('institution_type')['tier3_pubs'].mean()
overall_acc    = period_df['acceptance_rate'].mean()
resub_eval_avg = period_df['n_resub_evaluated'].mean()
resub_acc_avg  = period_df['n_resub_accepted'].mean()

add_stat_table(doc,
    headers=['Institution', 'Final AI Use', 'Total Pubs', 'Tier 1', 'Tier 2', 'Tier 3', 'Tenure Target'],
    rows=[
        ['R1',       f'{mean_ai_final["R1"]:.3f}',       f'{mean_pubs["R1"]:.1f}',
         f'{mean_t1["R1"]:.1f}', f'{mean_t2["R1"]:.1f}', f'{mean_t3["R1"]:.1f}', '10 (target 8–12)'],
        ['R2',       f'{mean_ai_final["R2"]:.3f}',       f'{mean_pubs["R2"]:.1f}',
         f'{mean_t1["R2"]:.1f}', f'{mean_t2["R2"]:.1f}', f'{mean_t3["R2"]:.1f}', '9 (target 6–12)'],
        ['Balanced', f'{mean_ai_final["Balanced"]:.3f}', f'{mean_pubs["Balanced"]:.1f}',
         f'{mean_t1["Balanced"]:.1f}', f'{mean_t2["Balanced"]:.1f}', f'{mean_t3["Balanced"]:.1f}', '6 (target 4–8)'],
        ['Teaching', f'{mean_ai_final["Teaching"]:.3f}', f'{mean_pubs["Teaching"]:.1f}',
         f'{mean_t1["Teaching"]:.1f}', f'{mean_t2["Teaching"]:.1f}', f'{mean_t3["Teaching"]:.1f}', '2 (target 1–3)'],
    ]
)

add_body(doc,
    f'Overall new-paper acceptance rate (all periods): {overall_acc*100:.1f}%. '
    f'Mean resubmissions evaluated per period: {resub_eval_avg:.0f}; '
    f'accepted per period: {resub_acc_avg:.0f} '
    f'({resub_acc_avg/(resub_eval_avg or 1)*100:.1f}% resubmission success rate).')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — Basic Descriptives
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Section 1: Basic Descriptives', level=1)

add_body(doc,
    'The simulation tracked 200 scholars across four institution types (R1, R2, Balanced, '
    'Teaching-Focused; 25% each) over 12 periods (6-year tenure clock). Scholars are '
    'initialized with AI use levels drawn from an empirical distribution (Nag et al., 2025). '
    'AI use evolves via pressure-driven reinforcement learning: each acceptance produces a '
    'positive step scaled by current publication pressure; each rejection produces a smaller '
    'negative step. Results below are averages across 50 replications; shaded bands indicate '
    '95% confidence intervals.')

add_heading(doc, 'AI Use Over Time', level=2, color=SECONDARY)
add_body(doc,
    f'AI use increased across all institution types over the 6-year tenure clock. '
    f'R1 scholars showed the steepest increase (final M = {mean_ai_final["R1"]:.2f}), '
    f'driven by high tenure pressure and relatively high acceptance rates at Tier 1 given '
    f'their research capacity advantage. Teaching-Focused scholars showed the slowest growth '
    f'(final M = {mean_ai_final["Teaching"]:.2f}), consistent with lower publication '
    f'pressure and fewer reinforcement signals from acceptance.')

add_figure(doc, 'outputs/report_figs/fig1_ai_use_time.png',
    'Figure 1a. Mean AI use level over time by institution type. '
    'Shaded bands = 95% CI across 50 replications.')

add_heading(doc, 'Acceptance Rate Over Time', level=2, color=SECONDARY)
add_body(doc,
    f'The overall new-paper acceptance rate averaged {overall_acc*100:.1f}% across all periods, '
    f'reflecting the percentile-threshold model (Tier 1: top 5%; Tier 2: top 15%; '
    f'Tier 3: top 30% of the quality distribution). The resubmission pipeline adds '
    f'a second pathway to acceptance: approximately {resub_eval_avg:.0f} resubmissions '
    f'are evaluated per period on average, of which {resub_acc_avg:.0f} ({resub_acc_avg/(resub_eval_avg or 1)*100:.1f}%) '
    f'succeed. The combined rate (new + resubmissions) is shown as a dashed line in '
    f'Figure 1b and is consistently higher than the new-paper-only rate.')

add_figure(doc, 'outputs/report_figs/fig1_acceptance_time.png',
    'Figure 1b. Overall acceptance rate over time. Solid = new submissions only; '
    'dashed = all evaluations (new + resubmissions). Shaded bands = 95% CI.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Efficiency
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Section 2: Efficiency', level=1)

add_body(doc,
    'Efficiency is defined as accepted publications per period. Scholars who use AI more '
    'extensively produce more papers per period (productivity multiplier: up to 2× at '
    'AI use = 1.0) but also produce lower-quality papers, which reduces per-paper '
    'acceptance rates. The net effect on accepted publications depends on the balance '
    'between volume and quality. Figures 2a and 2b show these competing dynamics.')

add_heading(doc, 'Publications per Period vs. AI Use', level=2, color=SECONDARY)
add_body(doc,
    'Figure 2a plots mean accepted publications per period against final AI use level, '
    'binned into deciles. R1 scholars exhibit the clearest positive relationship between '
    'AI use and output: high research capacity means that even at elevated AI use levels, '
    'paper quality remains sufficient for Tier 1 acceptance. Lower-capacity institutions '
    'show a flatter or potentially inverted relationship, as the quality penalty at high '
    'AI use offsets the volume gain. The resubmission pipeline partially cushions this '
    'effect by giving downgraded papers additional chances at lower tiers.')

add_figure(doc, 'outputs/report_figs/fig2_efficiency.png',
    'Figure 2a. Mean accepted publications per period vs. final AI use level (binned deciles), '
    'by institution type. Error bands = 95% CI.')

add_heading(doc, 'Acceptance Rate vs. AI Use', level=2, color=SECONDARY)
add_body(doc,
    'Figure 2b shows the per-scholar acceptance rate (accepted / total produced) as a '
    'function of final AI use level. A declining trend at high AI use levels reflects '
    'the quality penalty: scholars who rely heavily on AI produce more papers but at '
    'lower quality, depressing their hit rate. R1 scholars maintain higher acceptance '
    'rates across all AI use levels due to higher baseline research capacity.')

add_figure(doc, 'outputs/report_figs/fig2_accept_by_ai.png',
    'Figure 2b. Per-scholar acceptance rate vs. final AI use level, by institution type.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — Quality
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Section 3: Paper Quality', level=1)

add_body(doc,
    'Paper quality is computed via a piecewise penalty function calibrated to '
    'Gartenberg et al. (2026): quality = research_capacity − penalty(ai_use_level). '
    'The penalty is linear below AI_LOW_THRESHOLD (0.30) and accelerates above it '
    '(exponent = 1.50), reflecting the approximately 1.28 SD quality decline at '
    'high AI use documented in the empirical literature. Acceptance thresholds are '
    'derived from the assumed quality distribution Normal(0.50, 0.20): Tier 1 ≈ 0.829, '
    'Tier 2 ≈ 0.707, Tier 3 ≈ 0.605.')

add_heading(doc, 'Quality Over Time', level=2, color=SECONDARY)
add_body(doc,
    'Mean new-paper quality declines gradually over time as scholars increase AI use, '
    'with the steepest decline occurring in R1 (where AI adoption is fastest). '
    'Despite the decline, R1 and R2 scholars maintain quality well above the Tier 1 '
    'acceptance threshold (0.829) on average, reflecting their high research capacity. '
    'Balanced and Teaching-Focused scholars operate near or below the Tier 2 threshold, '
    'consistent with their quality-based tier targeting.')

add_figure(doc, 'outputs/report_figs/fig3_quality_time.png',
    'Figure 3a. Mean new-paper quality over time by institution type. '
    'Dashed reference lines show Tier 1 (0.829) and Tier 2 (0.707) acceptance thresholds.')

add_heading(doc, 'Quality vs. AI Use', level=2, color=SECONDARY)
add_body(doc,
    'Figure 3b contrasts quality for all papers (solid lines) against accepted papers only '
    '(dashed lines) as a function of final AI use level. The accepted-paper quality is '
    'uniformly higher than all-paper quality (as expected — accepted papers must exceed '
    'journal thresholds). The gap between all-paper and accepted-paper quality widens at '
    'high AI use levels, reflecting the increasing proportion of low-quality submissions '
    'that are consistently rejected (and eventually downgraded or abandoned via the '
    'resubmission pipeline).')

add_figure(doc, 'outputs/report_figs/fig3_quality_vs_ai.png',
    'Figure 3b. Mean paper quality vs. final AI use level. '
    'Solid = all papers produced; dashed = accepted papers only.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — Tenure Target Performance
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Section 4: Tenure Target Performance', level=1)

add_body(doc,
    'Tenure targets vary by institution type (R1: 8–12 total, with 3–6 Tier 1; '
    'R2: 6–12 total, 1–3 Tier 1; Balanced: 4–8 total, 2–4 Tier 2; Teaching: 1–3 total). '
    'The resubmission pipeline meaningfully increases total publication counts relative '
    'to the no-resubmission baseline, particularly for lower-tier institutions where '
    'papers cascade from Tier 2 to Tier 3 and accumulate acceptances.')

add_heading(doc, 'Publications vs. Tenure Targets', level=2, color=SECONDARY)
add_body(doc,
    f'Figure 4a shows mean total publications (including resubmission wins) at end of '
    f'the 6-year clock. R1 scholars exceed the tenure target midpoint substantially '
    f'(M = {mean_pubs["R1"]:.1f} vs. target = 10), driven by high capacity, frequent T1 '
    f'acceptance, and AI-driven volume. R2 scholars fall below target '
    f'(M = {mean_pubs["R2"]:.1f} vs. target = 9); their Tier 1 acceptance rate is '
    f'lower than R1\'s due to lower research capacity, and the resubmission pipeline '
    f'helps but does not close the gap. Balanced scholars (M = {mean_pubs["Balanced"]:.1f} '
    f'vs. target = 6) and Teaching-Focused scholars (M = {mean_pubs["Teaching"]:.1f} vs. '
    f'target = 2) are below or near target, respectively.')

add_figure(doc, 'outputs/report_figs/fig4_tenure_pubs.png',
    'Figure 4a. Mean total publications by institution type at end of tenure clock. '
    'Dashed lines = tenure target midpoints. Error bars = 95% CI.')

add_heading(doc, 'Tier Breakdown', level=2, color=SECONDARY)

t1_r1 = mean_t1['R1']
t1_r2 = mean_t1['R2']
add_body(doc,
    f'Figure 4b shows the tier composition of publications. R1 scholars average '
    f'{t1_r1:.1f} Tier 1 publications, closely matching the target midpoint of 4. '
    f'R2 scholars average only {t1_r2:.1f} Tier 1 publications (target midpoint = 2), '
    f'reflecting the challenge of clearing the 0.829 Tier 1 quality threshold with '
    f'lower research capacity, especially as AI use increases. The resubmission pipeline '
    f'converts many rejected T1 papers into T2 or T3 acceptances for R2 scholars, '
    f'boosting total count but not Tier 1 specifically. Balanced scholars '
    f'accumulate primarily Tier 2 and 3 publications, consistent with their '
    f'tier-targeting strategy (T1 thresholds set to 9.99).')

add_figure(doc, 'outputs/report_figs/fig4_tier_breakdown.png',
    'Figure 4b. Mean publications by tier and institution type. '
    'Dashed lines = Tier 1 publication targets (R1 and R2 only).')

add_heading(doc, 'Tenure Achievement by AI Use', level=2, color=SECONDARY)
add_body(doc,
    'Figure 4c shows the percentage of scholars meeting their total publication target '
    'as a function of AI use quartile. For R1 scholars, higher AI use is associated with '
    'better tenure target achievement (volume effect dominates). For Balanced and '
    'Teaching-Focused scholars, the relationship is flatter or modestly positive, '
    'as they operate in lower tiers where the volume benefit still exceeds the quality penalty.')

add_figure(doc, 'outputs/report_figs/fig4_tenure_pct_ai.png',
    'Figure 4c. Percentage of scholars meeting tenure publication targets '
    'by AI use quartile and institution type.')

add_heading(doc, 'Tier-Specific Target Achievement by AI Use', level=2, color=SECONDARY)
add_body(doc,
    'Figure 4d focuses on tier-specific targets: R1 and R2 scholars\' Tier 1 targets, '
    'and Balanced scholars\' Tier 2 target. For R1, achieving the Tier 1 target is '
    'slightly easier at lower-to-moderate AI use — where quality remains high enough '
    'to clear the T1 acceptance threshold. At very high AI use, quality degradation '
    'reduces T1 acceptance probability, working against the Tier 1 target even as '
    'total volume increases. This tension is the core dynamic the model is designed '
    'to capture.')

add_figure(doc, 'outputs/report_figs/fig4_tier_target_by_ai.png',
    'Figure 4d. Percentage of scholars meeting tier-specific targets '
    '(Tier 1 for R1/R2; Tier 2 for Balanced) by AI use quartile.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# APPENDIX — Resubmission Pipeline Notes
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Appendix: Resubmission Pipeline Details', level=1)
add_body(doc,
    'The resubmission pipeline was introduced in Version 3 of the model. Key features:')
bullets = [
    'Rejected papers enter the submitting scholar\'s resubmission_queue at the end '
    'of the period in which they were first evaluated (tier_attempts = 1).',
    'Each subsequent period, queued papers are evaluated first (before new paper '
    'production). The same paper quality is used, but a new Gaussian noise draw '
    '(SD = 0.10) is applied, capturing a fresh reviewer assessment.',
    'After MAX_TIER_ATTEMPTS (= 3) successive rejections at one tier, the paper '
    'downgrades to the next tier (journal_tier += 1; tier_attempts reset to 0). '
    'If the paper has already been rejected 3 times at Tier 3, it is abandoned.',
    'Accepted resubmissions are recorded in the scholar\'s publication_record and '
    'count toward total_publications, tier1_publications, tier2_publications, '
    'and the RL signal. They do NOT count toward papers_per_period (efficiency metric).',
    'Worst-case paper journey: 3 T1 rejections + 3 T2 rejections + 3 T3 rejections '
    '= 9 additional periods before abandonment. Given a 12-period simulation, papers '
    'produced in early periods can cycle through the full pipeline.',
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

doc.add_paragraph()
add_body(doc,
    f'In the current run, an average of {resub_eval_avg:.0f} resubmission evaluations '
    f'were performed per period across 200 scholars, compared to approximately '
    f'{period_df["n_produced"].mean():.0f} new paper evaluations per period. '
    f'The resubmission acceptance rate ({resub_acc_avg/(resub_eval_avg or 1)*100:.1f}%) '
    f'is lower than the new-paper acceptance rate ({overall_acc*100:.1f}%), reflecting '
    f'that papers in the resubmission queue have already been rejected at least once and '
    f'tend to be lower-quality submissions.')

doc.save(DOC_PATH)
print(f'\nSaved {DOC_PATH}')
print('Done.')
